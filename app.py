import os
import uuid
import subprocess
from datetime import datetime
from flask import Flask, request, send_file, render_template, redirect, flash, session, url_for
from werkzeug.utils import secure_filename
from PyPDF2 import PdfMerger
from docx import Document
from PIL import Image
from num2words import num2words
import shutil

# === CONFIG ===
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UPLOAD_FOLDER = os.path.join(BASE_DIR, "uploads")
TEMPLATE_DOCX = os.path.join(BASE_DIR, "EAF_Template.docx")
ALLOWED_EXTENSIONS = {"pdf", "png", "jpg", "jpeg"}
MAX_CONTENT_LENGTH = 200 * 1024 * 1024  # 200 MB

# Flask setup
app = Flask(__name__)
app.secret_key = os.environ.get("FLASK_SECRET_KEY", "dev-secret-change-this")
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = MAX_CONTENT_LENGTH
os.makedirs(UPLOAD_FOLDER, exist_ok=True)


app.config['SESSION_PERMANENT'] = False

# Simple credentials (override with env vars in production!)
VALID_USERS = {
    os.environ.get("APP_USERNAME", "admin"): os.environ.get("APP_PASSWORD", "password123")
}


# === UTILITIES ===
def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def number_to_words(amount):
    try:
        amt_int = int(float(amount))
    except Exception:
        return ""
    try:
        words = num2words(amt_int, lang="en_IN")
    except Exception:
        words = num2words(amt_int)
    return "Rupees " + words.replace("-", " ").title() + " Only"


def replace_placeholders_in_docx(doc: Document, mapping: dict):
    def replace_in_paragraph(paragraph):
        if not paragraph.text:
            return
        new_text = paragraph.text
        for key, val in mapping.items():
            if key in new_text:
                new_text = new_text.replace(key, val)
        if new_text != paragraph.text:
            # Clear all runs
            for run in paragraph.runs:
                run.text = ""
            # Put replaced text in the first run
            paragraph.runs[0].text = new_text

    # Replace in normal paragraphs
    for para in doc.paragraphs:
        replace_in_paragraph(para)

    # Replace inside tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para)

    # Replace in headers and footers
    for section in doc.sections:
        for para in section.header.paragraphs:
            replace_in_paragraph(para)
        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para)

        for para in section.footer.paragraphs:
            replace_in_paragraph(para)
        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        replace_in_paragraph(para)


def generate_eaf_docx(template_path, out_docx_path, date_str, amount, amount_words, event_name, remarks, budget_name, budget_head,
                      budgeted_amount, amount_spent, balance_available, payment_type, acc_number, acc_holder, bank_name, ifsc, branch):
    vendor_payment = "✔" if payment_type == "vendor_payment" else ""
    reimbursement = "✔" if payment_type == "reimbursement" else ""
    advance = "✔" if payment_type == "advance" else ""
    
    doc = Document(template_path)
    mapping = {
        "{{DATE}}": date_str,
        "{{TOTAL_AMOUNT}}": str(amount),
        "{{TOTAL_AMOUNT_WORDS}}": amount_words,
        "{{EVENT_NAME}}": event_name,
        "{{REMARKS}}": remarks,
        "{{BUDGET_NAME}}": budget_name,
        "{{BUDGET_HEAD}}": budget_head,
        "{{BUDGETED_AMOUNT}}": budgeted_amount,
        "{{AMOUNT_SPENT}}": amount_spent,
        "{{BALANCE_AVAILABLE}}": balance_available,
        "{{V}}": vendor_payment,
        "{{R}}": reimbursement,
        "{{A}}": advance,
        "{{ACCOUNT_NUMBER}}": acc_number,
        "{{ACCOUNT_HOLDER}}": acc_holder,
        "{{BANK_NAME}}": bank_name,
        "{{IFSC}}": ifsc,
        "{{BRANCH}}": branch
    }
    replace_placeholders_in_docx(doc, mapping)
    doc.save(out_docx_path)
    return out_docx_path


def convert_docx_to_pdf(docx_path, pdf_path):
    outdir = os.path.dirname(pdf_path)
    possible_paths = [
        "/usr/bin/libreoffice",
        "/usr/bin/soffice",
        shutil.which("libreoffice"),
        shutil.which("soffice")
    ]
    soffice_path = next((p for p in possible_paths if p and os.path.exists(p)), None)
    if not soffice_path:
        raise RuntimeError("LibreOffice not found. Install it in Dockerfile or system.")

    subprocess.run(
        [soffice_path, "--headless", "--convert-to", "pdf:writer_pdf_Export", "--outdir", outdir, docx_path],
        check=True
    )

    generated_pdf = os.path.join(outdir, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    if generated_pdf != pdf_path:
        os.replace(generated_pdf, pdf_path)

    if not os.path.exists(pdf_path):
        raise RuntimeError("PDF conversion failed")

    return pdf_path


# === ROUTES ===
@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username")
        password = request.form.get("password")

        if username in VALID_USERS and VALID_USERS[username] == password:
            session["user"] = username
            flash("Login successful!")
            return redirect(url_for("index"))
        else:
            flash("Invalid username or password.")
            return redirect(url_for("login"))

    return render_template("login.html")


@app.route("/", methods=["GET", "POST"])
def index():
    if "user" not in session:
        return redirect(url_for("login"))

    today = datetime.now().strftime("%Y-%m-%d")  # format for <input type="date">

    if request.method == "POST":
        # Form data
        date = request.form.get("date") or today
        date = datetime.strptime(date, "%Y-%m-%d").strftime("%d-%m-%Y")
        amount = request.form.get("amount") or "0"
        amount_words = request.form.get("amount_words") or number_to_words(amount)
        event_name = request.form.get("event_name") or ""
        budget_name = request.form.get("budget_name") or ""
        budget_head = request.form.get("budget_head") or ""
        budgeted_amount = request.form.get("budgeted_amount") or ""
        amount_spent = request.form.get("amount_spent") or ""
        balance_available = ""
        if (budgeted_amount != "" and amount_spent != ""):
            balance_available = str(int(budgeted_amount) - int(amount_spent))
        payment_type = request.form.get("payment_type")
        purpose = request.form.get("purpose") or ""
        bundle_filename = request.form.get("filename") or f"Bundle_{datetime.now().strftime('%d%m%Y_%H%M%S')}"
        bills_only_filename = bundle_filename + "_Bill"
        # Optional Bank details

        acc_number = request.form.get("account_number") or ""
        acc_holder = request.form.get("account_holder") or ""
        bank_name = request.form.get("bank_name") or ""
        ifsc = request.form.get("ifsc") or ""
        branch = request.form.get("branch") or ""

        bill_files = request.files.getlist("bills")
        if not bill_files or all(f.filename == "" for f in bill_files):
            flash("Please upload at least one bill file (pdf or image).")
            return redirect(request.url)

        saved_pdf_paths = []
        for f in bill_files:
            if f and allowed_file(f.filename):
                safe_name = secure_filename(f.filename)
                unique_name = f"{uuid.uuid4().hex}_{safe_name}"
                save_path = os.path.join(app.config["UPLOAD_FOLDER"], unique_name)
                f.save(save_path)
                ext = safe_name.rsplit(".", 1)[1].lower()
                if ext in ("png", "jpg", "jpeg"):
                    img = Image.open(save_path)
                    if img.mode in ("RGBA", "LA", "P"):
                        img = img.convert("RGB")
                    pdf_path = save_path + ".pdf"
                    img.save(pdf_path, "PDF", resolution=100.0)
                    os.remove(save_path)
                    saved_pdf_paths.append(pdf_path)
                else:
                    saved_pdf_paths.append(save_path)
            else:
                flash(f"File '{f.filename}' not allowed. Allowed: pdf, png, jpg, jpeg")
                return redirect(request.url)

        # Generate EAF DOCX
        timestamp = datetime.now().strftime("%d%m%Y_%H%M%S")
        out_docx = os.path.join(app.config["UPLOAD_FOLDER"], f"EAF_{timestamp}.docx")
        generate_eaf_docx(TEMPLATE_DOCX, out_docx, date, amount, amount_words, event_name, purpose, budget_name, budget_head,
                          budgeted_amount, amount_spent, balance_available, payment_type, acc_number, acc_holder, bank_name, ifsc, branch)

        # Convert DOCX -> PDF
        out_pdf = os.path.join(app.config["UPLOAD_FOLDER"], f"EAF_{timestamp}.pdf")
        try:
            convert_docx_to_pdf(out_docx, out_pdf)
        except Exception as e:
            flash(f"Failed to convert DOCX to PDF: {e}")
            return redirect(request.url)

        # Merge EAF PDF + bills
        merged_pdf_path = os.path.join(app.config["UPLOAD_FOLDER"], f"{bundle_filename}.pdf")
        try:
            with PdfMerger() as merger:
                merger.append(out_pdf)
                for p in saved_pdf_paths:
                    merger.append(p)
                merger.write(merged_pdf_path)
        except Exception as me:
            flash("Error merging PDFs: " + str(me))
            return redirect(request.url)

        # Bills only PDF
        bills_only_pdf_path = os.path.join(app.config["UPLOAD_FOLDER"], f"{bills_only_filename}.pdf")
        try:
            with PdfMerger() as merger2:
                for p in saved_pdf_paths:
                    merger2.append(p)
                merger2.write(bills_only_pdf_path)
        except Exception as me:
            flash("Error creating bills-only PDF: " + str(me))
            return redirect(request.url)

        return render_template("download.html",
                               bundle_pdf=os.path.basename(merged_pdf_path),
                               bills_pdf=os.path.basename(bills_only_pdf_path))

    return render_template("index.html", today = today)


@app.route("/uploads/<filename>")
def uploaded_file(filename):
    if "user" not in session:
        flash("Please log in first.")
        return redirect(url_for("login"))

    file_path = os.path.join(app.config["UPLOAD_FOLDER"], filename)
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        flash("File not found.")
        return redirect("/")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("login"))


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.environ.get("PORT", 5000)))
